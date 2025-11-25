#!/usr/bin/env python3
"""
Recursive document scanner with text extraction, lightweight classification, and rich JSON output.
Supports UNC paths, include/exclude filters, allowed extensions, metadata capture,
language detection, and rule-based tagging.
"""

from __future__ import annotations

import argparse
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Set, Tuple

from file_summary import summarize_text

try:
    import yaml  # type: ignore
except ImportError:  # Optional dependency
    yaml = None

try:
    import docx  # type: ignore
except ImportError:
    docx = None

try:
    import pdfplumber  # type: ignore
except ImportError:
    pdfplumber = None

try:
    import PyPDF2  # type: ignore
except ImportError:
    PyPDF2 = None

try:
    from langdetect import detect, LangDetectException  # type: ignore
except ImportError:
    detect = None
    LangDetectException = Exception

EXCLUDED_FILE_PATTERNS = (
    "Thumbs.db",
    "~$",
)


def is_excluded_file(path: Path) -> bool:
    """Return True if file is a temporary, hidden, or unwanted file."""
    name = path.name.lower()

    # Exact matches
    if name == "thumbs.db":
        return True

    # Word temporary files
    if name.startswith("~$"):
        return True

    # Generic garbage
    if name.endswith(".tmp"):
        return True

    return False

# ---------- Logging ----------
def log(msg: str) -> None:
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{timestamp}] {msg}")


# ---------- Config handling ----------
def read_config(config_path: Path) -> Dict[str, Any]:
    if not config_path.exists():
        raise FileNotFoundError(f"Config not found: {config_path}")

    suffix = config_path.suffix.lower()
    raw_text = config_path.read_text(encoding="utf-8", errors="ignore")

    if suffix in {".yaml", ".yml"}:
        if yaml is None:
            raise RuntimeError("pyyaml is required for YAML configs.")
        data = yaml.safe_load(raw_text)
    else:
        data = json.loads(raw_text)

    if not isinstance(data, dict):
        raise ValueError("Config must define a JSON/YAML object.")
    return data


def normalize_paths(entries: Iterable[str]) -> List[Path]:
    normalized: List[Path] = []
    for entry in entries:
        entry = entry.strip()
        if entry:
            normalized.append(Path(entry.strip("\\/")))
    return normalized


def normalize_excludes(entries: Iterable[str]) -> List[Tuple[str, ...]]:
    """Return excluded folders as tuples of path parts for prefix matching."""
    paths: List[Tuple[str, ...]] = []
    for p in normalize_paths(entries):
        if p.parts:
            paths.append(p.parts)
    return paths


def normalize_extensions(entries: Iterable[str]) -> Set[str]:
    return {ext.lower().lstrip(".") for ext in entries if ext}


# ---------- Text extraction helpers ----------
def read_text_file(path: Path, max_chars: int = 20000) -> Tuple[str, int, bool, int]:
    """Return text, page_count, contains_images, tables_count."""
    try:
        text = path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        return f"Unable to read file: {exc}", 1, False, 0
    text = text[:max_chars]
    return text, 1, False, 0


def extract_docx(path: Path, max_chars: int = 20000) -> Tuple[str, int, bool, int]:
    if docx is None:
        return "python-docx not installed.", 0, False, 0
    try:
        document = docx.Document(path)
    except Exception as exc:  # noqa: BLE001
        return f"Unable to read DOCX: {exc}", 0, False, 0

    paragraphs = [p.text for p in document.paragraphs if p.text]
    text = "\n".join(paragraphs)
    tables_count = len(document.tables)

    contains_images = False
    try:
        rels = getattr(document.part, "related_parts", {})
        contains_images = any(
            getattr(part, "content_type", "").startswith("image/") for part in rels.values()
        )
    except Exception:  # noqa: BLE001
        contains_images = False

    return text[:max_chars], max(1, len(paragraphs) // 40), contains_images, tables_count


def extract_pdf(path: Path, max_chars: int = 20000) -> Tuple[str, int, bool, int]:
    if pdfplumber:
        try:
            with pdfplumber.open(path) as pdf:
                texts: List[str] = []
                contains_images = False
                tables_count = 0
                for page in pdf.pages:
                    txt = page.extract_text() or ""
                    texts.append(txt)
                    try:
                        if getattr(page, "images", []):
                            contains_images = True
                    except Exception:  # noqa: BLE001
                        pass
                    try:
                        tables = page.extract_tables()
                        tables_count += len(tables)
                    except Exception:  # noqa: BLE001
                        pass
                joined = "\n".join(texts)[:max_chars]
                return joined, len(pdf.pages), contains_images, tables_count
        except Exception as exc:  # noqa: BLE001
            return f"Unable to read PDF: {exc}", 0, False, 0

    if PyPDF2:
        try:
            reader = PyPDF2.PdfReader(str(path))
            texts: List[str] = []
            for page in reader.pages:
                try:
                    texts.append(page.extract_text() or "")
                except Exception:  # noqa: BLE001
                    continue
            joined = "\n".join(texts)[:max_chars]
            return joined, len(reader.pages), False, 0
        except Exception as exc:  # noqa: BLE001
            return f"Unable to read PDF: {exc}", 0, False, 0

    return "pdfplumber or PyPDF2 not installed.", 0, False, 0


def extract_text(path: Path) -> Tuple[str, int, bool, int]:
    ext = path.suffix.lower().lstrip(".")
    if ext in {"txt", "md", "log"}:
        return read_text_file(path)
    if ext in {"docx"}:
        return extract_docx(path)
    if ext in {"pdf"}:
        return extract_pdf(path)
    return "Unsupported extension for extraction.", 0, False, 0


# ---------- Classification helpers ----------
MODULE_KEYWORDS = {
    "MES": ["mes", "manufacturing execution"],
    "APS": ["aps", "advanced planning"],
    "SFC": ["sfc", "shop floor"],
}

PROCESS_STEPS = {
    "analisi": ["analysis", "analisi", "requirements"],
    "sviluppo": ["development", "sviluppo", "implementazione"],
    "test": ["test", "testing", "collaudo", "validation"],
    "produzione": ["production", "go-live", "deploy", "rilascio"],
}

CONTENT_TYPES = {
    "manual": ["manual", "user guide", "istruzioni"],
    "specification": ["specification", "specifica", "requirement"],
    "test case": ["test case", "piano di test", "caso di test"],
    "design": ["design", "architettura"],
}


def detect_language(text: str) -> str:
    cleaned = text[:2000]
    if not cleaned.strip():
        return "unknown"
    if detect:
        try:
            return detect(cleaned)
        except LangDetectException:
            return "unknown"
    # Heuristic fallback: count vowels typical of Italian vs English
    it_score = len(re.findall(r"\b(il|la|lo|gli|che|per|con|una|un)\b", cleaned.lower()))
    en_score = len(re.findall(r"\b(the|and|for|with|this|that|of)\b", cleaned.lower()))
    if it_score > en_score:
        return "it"
    if en_score > it_score:
        return "en"
    return "unknown"


def detect_modules(text: str) -> List[str]:
    hits = []
    lower = text.lower()
    for module, keywords in MODULE_KEYWORDS.items():
        if any(k in lower for k in keywords):
            hits.append(module)
    return hits


def detect_process_step(text: str) -> str:
    lower = text.lower()
    for step, keywords in PROCESS_STEPS.items():
        if any(k in lower for k in keywords):
            return step
    return ""


def detect_content_type(text: str) -> str:
    lower = text.lower()
    for ctype, keywords in CONTENT_TYPES.items():
        if any(k in lower for k in keywords):
            return ctype
    return ""


def detect_version(text: str) -> str:
    matches = re.findall(r"\b(R\\d{2}|V\\d+(?:\\.\\d+)?|Draft)\\b", text, flags=re.IGNORECASE)
    return matches[0] if matches else ""


def compute_complexity(word_count: int) -> int:
    if word_count < 200:
        return 1
    if word_count < 800:
        return 2
    if word_count < 1500:
        return 3
    if word_count < 3000:
        return 4
    return 5


# ---------- Scanning ----------
def should_skip_folder(current_parts: Sequence[str], excluded: List[Tuple[str, ...]]) -> bool:
    """Return True if current relative path contains any excluded path parts."""
    for parts in excluded:
        if not parts:
            continue
        for idx in range(0, len(current_parts) - len(parts) + 1):
            if tuple(current_parts[idx : idx + len(parts)]) == parts:
                return True
    return False


def gather_files(
    root: Path,
    includes: Iterable[Path],
    excludes: List[Tuple[str, ...]],
    allowed_exts: Set[str],
) -> List[Path]:

    files: List[Path] = []

    for include in includes:
        include_root = (root / include).resolve()

        if not include_root.exists():
            log(f"Include path missing, skipped: {include_root}")
            continue

        # Calcola la parte relativa rispetto a root
        try:
            rel_parts = include_root.relative_to(root).parts
        except ValueError:
            # include_root NON è sotto root → skip
            log(f"Include path outside root, skipped: {include_root}")
            continue

        # Se la cartella include è dentro una exclude_folder → skip
        if should_skip_folder(rel_parts, excludes):
            log(f"Include path excluded by filter, skipped: {include_root}")
            continue

        # Legge SOLO i file DIRETTI nella cartella
        try:
            for entry in include_root.iterdir():
                if entry.is_file():
                    # Auto-exclude garbage files
                    if is_excluded_file(entry):
                        log(f"Skipping temporary/hidden file: {entry}")
                        continue

                    ext = entry.suffix.lower().lstrip(".")
                    if allowed_exts and ext not in allowed_exts:
                        continue
                    files.append(entry)
        except OSError as exc:
            log(f"Unable to access include path {include_root}: {exc}")

    return files



# ---------- Main processing ----------
def process_file(path: Path, tag: str, domain: str) -> Dict[str, Any]:
    text, page_count, contains_images, tables_count = extract_text(path)
    cleaned_summary = summarize_text(text, max_len=600)
    language = detect_language(text)
    word_count = len(cleaned_summary.split())
    complexity = compute_complexity(word_count)
    modules_mentioned = detect_modules(text)
    process_step = detect_process_step(text)
    content_type = detect_content_type(text)
    version = detect_version(text)
    is_confidential = bool(re.search(r"\bconfidential|riservato|internal\b", text, re.IGNORECASE))

    stat = path.stat()
    size_kb = round(stat.st_size / 1024, 2)
    created_at = datetime.fromtimestamp(stat.st_ctime).isoformat()
    modified_at = datetime.fromtimestamp(stat.st_mtime).isoformat()

    return {
        "path": str(path),
        "filename": path.name,
        "tag": tag,
        "summary": cleaned_summary,
        "extension": path.suffix.lower(),
        "size_kb": size_kb,
        "created_at": created_at,
        "modified_at": modified_at,
        "language": language,
        "page_count": page_count,
        "word_count": word_count,
        "content_type": content_type,
        "domain": domain,
        "customers_mentioned": [],  # placeholder for future NLP
        "modules_mentioned": modules_mentioned,
        "process_step": process_step,
        "version": version,
        "contains_images": contains_images,
        "tables_count": tables_count,
        "complexity": complexity,
        "is_confidential": is_confidential,
    }


def build_output(
    root: Path,
    files: List[Path],
    tag: str,
    domain: str,
) -> Dict[str, Any]:
    cliente = root.name or root.parts[-1] if root.parts else "unknown"
    output_files = []
    for path in files:
        #DBG log
        #log(f"Processing file: {path}")
        output_files.append(process_file(path, tag=tag, domain=domain))
    return {"cliente": cliente, "count": len(output_files), "files": output_files}


def save_output(data: Dict[str, Any], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    log(f"Output saved to: {output_path}")


# ---------- CLI ----------
def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Scan documents and summarize them.")
    parser.add_argument(
        "-c",
        "--config",
        type=Path,
        default=Path("docs_config.json"),
        help="Path to config file (JSON or YAML).",
    )
    parser.add_argument(
        "-o",
        "--output-dir",
        type=Path,
        default=None,
        help="Directory where output.json will be saved (default: alongside script).",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    script_dir = Path(__file__).resolve().parent
    config_path = (args.config if args.config.is_absolute() else script_dir / args.config).resolve()

    try:
        config = read_config(config_path)
    except Exception as exc:  # noqa: BLE001
        log(f"Failed to read config: {exc}")
        return

    root = Path(str(config.get("input_root_path", ""))).expanduser()
    if not root.exists():
        log(f"Root path not found or inaccessible: {root}")
        return

    include_paths = normalize_paths(config.get("include_paths", config.get("include_folders", [])))
    exclude_paths = normalize_excludes(
        config.get("exclude_paths", config.get("exclude_folders", []))
    )
    allowed_exts = normalize_extensions(config.get("allowed_extensions", []))
    tag = str(config.get("tag", "")).strip()
    domain = str(config.get("domain", "")).strip()

    log(f"Scanning root: {root}")
    log(f"Including: {', '.join(str(p) for p in include_paths) if include_paths else '(none)'}")
    log(
        f"Excluding: {', '.join('/'.join(parts) for parts in exclude_paths) if exclude_paths else '(none)'}"
    )
    log(f"Allowed extensions: {', '.join(sorted(allowed_exts)) if allowed_exts else '(all)'}")

    try:
        files = gather_files(root, include_paths or [Path(".")], exclude_paths, allowed_exts)
    except Exception as exc:  # noqa: BLE001
        log(f"Error during scan: {exc}")
        return

    output_data = build_output(root, files, tag=tag, domain=domain)

    output_dir = args.output_dir if args.output_dir is not None else script_dir
    output_path = output_dir / "output.json"

    try:
        save_output(output_data, output_path)
    except Exception as exc:  # noqa: BLE001
        log(f"Failed to save output: {exc}")
        return

    print(output_path)


if __name__ == "__main__":
    main()
